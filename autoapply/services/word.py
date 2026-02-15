import asyncio
import os
import logging

from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Optional

from autoapply.logging import get_logger
from autoapply.models import Certification, Contact, Education, JobExperience, Skills

get_logger()
logger = logging.getLogger(__name__)


def create_resume(
    save_path: str,
    contact: Contact,
    summary_text: str,
    job_exp: list[JobExperience],
    skills: list[Skills],
    education_entries: list[Education],
    certifications: list[Certification],
):
    doc = Document()

    # 1. Setup Margins (Narrow margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Helper function to format fonts
    def set_font(run, font_name="Arial", font_size=10, bold=False, italic=False):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic

    # Helper to add a bottom border
    def add_bottom_border(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # --- HEADER ---
    # [cite_start]Name [cite: 1]
    name_para = doc.add_paragraph()
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(contact.name)
    set_font(name_run, font_size=10, bold=True)

    # [cite_start]Contact Info [cite: 2]
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    contact_para.paragraph_format.space_after = Pt(0)
    contact_text = f"{contact.location} | {contact.phone} | {contact.email} | {contact.linkedin} | {contact.github}"
    contact_run = contact_para.add_run(contact_text)
    set_font(contact_run, font_size=10)

    # --- SUMMARY ---
    # [cite_start]Heading [cite: 3]
    h_summary = doc.add_paragraph("Summary")
    h_summary.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h_summary.paragraph_format.space_after = Pt(0)
    h_summary_run = h_summary.runs[0]
    set_font(h_summary_run, font_size=10, bold=True)
    add_bottom_border(h_summary)

    # [cite_start]Content [cite: 4-6]
    p_summary = doc.add_paragraph(summary_text)
    p_summary.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_summary.paragraph_format.space_after = Pt(0)
    set_font(p_summary.runs[0], font_size=10)

    # --- WORK EXPERIENCE ---
    # [cite_start]Heading [cite: 7]
    h_exp = doc.add_paragraph("Work Experience")
    h_exp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h_exp.paragraph_format.space_after = Pt(0)
    h_exp_run = h_exp.runs[0]
    set_font(h_exp_run, font_size=10, bold=True)
    add_bottom_border(h_exp)

    # Helper for Job Entries (Table Layout)
    def add_job_entry(title_line, date_line, bullets):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        # --- FIX 1: Adjust Column Widths ---
        # Increased Left to 5.9" (was 5.75) to prevent location wrapping
        # Decreased Right to 1.6" (was 1.75) which is just enough for "Jan 2024 - Dec 2025"
        table.columns[0].width = Inches(5.9)
        table.columns[1].width = Inches(1.6)

        # Left Cell (Title | Company | Loc)
        cell_1 = table.cell(0, 0)
        cell_1.width = Inches(5.9)  # Explicit width for cell
        p1 = cell_1.paragraphs[0]
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(title_line)
        set_font(r1, font_size=10, bold=True)

        # Right Cell (Dates)
        cell_2 = table.cell(0, 1)
        cell_2.width = Inches(1.6)  # Explicit width for cell
        p2 = cell_2.paragraphs[0]
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(date_line)
        set_font(r2, font_size=10, bold=True)

        # Bullets
        for bullet in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(bullet)
            set_font(r, font_size=10)

    for job in job_exp:
        # --- FIX 2: Correct Quote Syntax in f-strings (Use ' inside "") ---
        if isinstance(job.to_, date):
            add_job_entry(
                f"{job.job_title} | {job.company_name} | {job.location}",
                f"{job.from_.strftime('%b %Y')} - {job.to_.strftime('%b %Y')}",
                job.experience,
            )
        else:
            add_job_entry(
                f"{job.job_title} | {job.company_name} | {job.location}",
                f"{job.from_.strftime('%b %Y')} - Current",
                job.experience,
            )

    # --- SKILLS ---
    # [cite_start]Heading [cite: 27]
    h_skills = doc.add_paragraph("Skills")
    h_skills.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h_skills.paragraph_format.space_after = Pt(0)
    h_skills_run = h_skills.runs[0]
    set_font(h_skills_run, font_size=10, bold=True)
    add_bottom_border(h_skills)

    # Helper for Skills
    def add_skill_line(category, items):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        r_cat = p.add_run(category + ": ")
        set_font(r_cat, font_size=10, bold=True)
        r_item = p.add_run(items)
        set_font(r_item, font_size=10)

    for skill in skills:
        add_skill_line(skill.title, skill.skills)

    # --- EDUCATION ---
    # [cite_start]Heading [cite: 33]
    h_edu = doc.add_paragraph("Education")
    h_edu.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h_edu.paragraph_format.space_after = Pt(0)
    h_edu_run = h_edu.runs[0]
    set_font(h_edu_run, font_size=10, bold=True)
    add_bottom_border(h_edu)

    # [cite_start]Education Entries [cite: 34]
    def add_education_section(left_cell: str, right_cell: str):
        table_edu1 = doc.add_table(rows=1, cols=2)
        table_edu1.autofit = False

        # --- FIX 3: Apply same width logic to Education ---
        table_edu1.columns[0].width = Inches(5.9)
        table_edu1.columns[1].width = Inches(1.6)

        cell_1 = table_edu1.cell(0, 0)
        cell_1.width = Inches(5.9)
        p1 = cell_1.paragraphs[0]
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(left_cell)
        set_font(r1, font_size=10, bold=False)

        cell_2 = table_edu1.cell(0, 1)
        cell_2.width = Inches(1.6)
        p2 = cell_2.paragraphs[0]
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(right_cell)
        set_font(r2, font_size=10, bold=False)

    for edu in education_entries:
        add_education_section(
            f"{edu.degree} | {edu.major} | {edu.college}",
            f"{edu.from_.strftime('%b %Y')} - {edu.to_.strftime('%b %Y')}",
        )

    # --- CERTIFICATIONS ---
    # [cite_start]Heading [cite: 35]
    h_cert = doc.add_paragraph("Certifications")
    h_cert.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h_cert.paragraph_format.space_after = Pt(0)
    h_cert_run = h_cert.runs[0]
    set_font(h_cert_run, font_size=10, bold=True)
    add_bottom_border(h_cert)

    # [cite_start]Content [cite: 36-37]
    # Note: Using the passed `certificate` string instead of hardcoded values
    def add_certification_section(certificate: str):
        p_cert1 = doc.add_paragraph(certificate)
        p_cert1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_cert1.paragraph_format.space_after = Pt(0)
        set_font(p_cert1.runs[0], font_size=10)

    for certificate in certifications:
        if certificate.expiry_date:
            add_certification_section(
                f"{certificate.title} | {certificate.obtained_date} | {certificate.expiry_date}"
            )
        else:
            add_certification_section(
                f"{certificate.title} | {certificate.obtained_date}"
            )

    # Save the file
    resume = os.path.join(save_path, "shashank_reddy.docx")
    doc.save(resume)
    print(f"Resume created successfully: {resume}")
    return resume


if __name__ == "__main__":
    contact = Contact(
        name="Shashank Shashishekhar Reddy",
        email="shnkreddy98@gmail.com",
        location="San Jose, California",
        phone="(510) 892-7191",
        linkedin="linkedin.com/in/shnkreddy",
        github="github.com/shnkreddy98",
    )
    summary = "Founding Data Engineer with 5+ years of experience architecting high-throughput, event-driven data platforms. Expert in building scalable ELT/ETL pipelines using AWS, Spark, Kafka, and ClickHouse, achieving sub-second latency. Proven track record of delivering 0-to-1 enterprise data solutions and driving 30%+ efficiency gains through automation."
    jobs = [
        JobExperience(
            job_title="Founding Engineer",
            company_name="AirFold",
            location="San Francisco, California",
            from_=date(year=2024, day=1, month=1),
            to_="current",
            experience=[
                "Architected multi-cloud platform (AWS/GCP) using Terraform, EKS, and Istio, ensuring high availability for 50+ services and enhancing system resilience.",
                "Engineered event-driven pipeline (Kafka, Debezium CDC, ClickHouse) processing millions of daily events with sub-second latency and exactly-once semantics.",
                "Developed 100+ production-grade FastAPI endpoints with robust Auth0/RBAC security, serving as the backbone for scalable client interactions.",
                "Integrated dbt into the ELT pipeline to standardize transformation logic in Snowflake & ClickHouse, reducing reporting errors by 40% via automated quality tests.",
                "Implemented scalable K8s infrastructure with Karpenter auto-scaling and FluxCD GitOps, automating continuous delivery for over 60 Helm releases.",
                "Managed multi-database architecture (PostgreSQL, DynamoDB, ClickHouse), optimizing storage for vector search, caching, and large-scale analytics.",
            ],
        ),
        JobExperience(
            job_title="Data Engineer",
            company_name="Kantar",
            location="Bengaluru, India",
            from_=date(year=2020, day=1, month=4),
            to_=date(year=2022, day=1, month=4),
            experience=[
                "Orchestrated ELT pipelines using Airflow and Spark, integrating Kafka streams to ingest advertising data with exactly-once semantics.",
                "Accelerated data modeling by deploying dbt for complex SQL transformations, replacing ad-hoc scripts with modular pipelines that reduced time-to-insight by 25%.",
                "Designed real-time processing framework with Kafka Streams and Spark Streaming, providing near real-time insights into campaign performance.",
                "Deployed serverless architecture (AWS Lambda, S3) to automate data workflows, reducing manual intervention by 80% and improving reliability.",
                "Engineered ROI optimization algorithms that increased client sales by 34% for global enterprises through data-driven ad spend allocation.",
            ],
        ),
        JobExperience(
            job_title="Data Engineer",
            company_name="The Sparks Foundation",
            location="Bengaluru, India",
            from_=date(year=2018, day=1, month=3),
            to_=date(year=2020, day=1, month=3),
            experience=[
                "Optimized ETL pipeline performance, achieving 25% efficiency improvement through parallel processing and exponential backoff retry logic.",
                "Enhanced PostgreSQL & MySQL query performance by 30% through strategic indexing and execution plan analysis.",
                "Integrated automated data validation scripts using Python and Pandas to detect schema mismatches and null values, reducing downstream data corruption incidents by 40%.",
                "Collaborated with business analysts to design star-schema data models in PostgreSQL, enabling faster generation of daily operational reports and supporting strategic decision-making.",
                "Mentored junior engineers on TDD and code reviews, establishing 80% test coverage targets and improving code maintainability.",
            ],
        ),
    ]
    skills = [
        Skills(title="Languages", skills="Python, SQL, Java, Go, Bash, TypeScript"),
        Skills(
            title="Data Engineering",
            skills="Apache Spark, PySpark, Apache Kafka, Airflow, dbt, Debezium CDC, Ray Serve, ETL/ELT",
        ),
        Skills(
            title="Cloud & Infra",
            skills="AWS (EC2, S3, Lambda, EKS, RDS), GCP, Terraform, Kubernetes, Docker, Istio, GitOps",
        ),
        Skills(
            title="Databases",
            skills="Snowflake, ClickHouse, PostgreSQL, DynamoDB, Redshift, MySQL, Redis, Delta Lake",
        ),
        Skills(
            title="Backend & Tools",
            skills="FastAPI, RESTful APIs, Git/GitHub Actions, Pytest, Pandas, NumPy",
        ),
    ]
    education = [
        Education(
            degree="Master's Degree",
            major="Data Analytics",
            college="San Jose State University",
            from_=date(year=2023, day=1, month=1),
            to_=date(year=2024, day=1, month=12),
        ),
        Education(
            degree="Post Graduate Diploma",
            major="Data Science",
            college="IIIT Bangalore",
            from_=date(year=2020, day=1, month=11),
            to_=date(year=2021, day=1, month=10),
        ),
        Education(
            degree="Bachelor of Engineering",
            major="Computer Science",
            college="Visvesvaraya Technological University",
            from_=date(year=2016, day=1, month=7),
            to_=date(year=2020, day=1, month=8),
        ),
    ]
    certificates = [
        "AWS Certified Cloud Practitioner | Amazon Web Services (AWS) | Jul 2023",
        "Azure AI Fundamentals | Microsoft Certified | Jan 2025",
    ]

    resume_name = create_resume(
        contact=contact,
        summary_text=summary,
        job_exp=jobs,
        skills=skills,
        education_entries=education,
        certifications=certificates,
    )


async def convert_docx_to_pdf(resume_docx: str) -> Optional[str]:
    """Convert DOCX to PDF using LibreOffice"""

    # Convert to absolute path
    docx_path = Path(resume_docx).resolve()

    # Verify file exists
    if not docx_path.exists():
        logger.error(f"DOCX file not found: {docx_path}")
        return None

    # Get output directory and expected PDF path
    output_dir = str(docx_path.parent)
    expected_pdf = str(docx_path.with_suffix(".pdf"))

    logger.debug(f"Converting: {docx_path}")
    logger.debug(f"Output dir: {output_dir}")
    logger.debug(f"Expected PDF: {expected_pdf}")

    try:
        # Create the subprocess to convert DOCX to PDF
        process = await asyncio.create_subprocess_exec(
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            str(docx_path),
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )

        # Wait for it to finish
        stdout, stderr = await process.communicate()

        if process.returncode == 0:
            # Verify the PDF was created
            if os.path.exists(expected_pdf):
                logger.debug(f"Successfully converted to PDF: {expected_pdf}")
                return expected_pdf
            else:
                logger.error(f"Conversion reported success but PDF not found: {expected_pdf}")
                return None
        else:
            error_msg = stderr.decode().strip()
            logger.error(f"LibreOffice conversion failed: {error_msg}")
            logger.error(f"Return code: {process.returncode}")
            return None

    except FileNotFoundError as e:
        logger.error(f"LibreOffice not found: {e}")
        logger.error("Make sure LibreOffice is installed in Docker")
        return None
    except Exception as e:
        logger.error(f"Exception during PDF conversion: {e}", exc_info=True)
        return None
