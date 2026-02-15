import asyncio
import base64
import uuid
import os
import subprocess
import logging
from typing import Dict, List, Any
from textwrap import dedent

from playwright.async_api import (
    Page,
)
from docx import Document

from autoapply.services.llm.models import (
    BrowserClickArgs,
    BrowserCloseArgs,
    BrowserConsoleMessagesArgs,
    BrowserDragArgs,
    BrowserEvaluateArgs,
    BrowserFileUploadArgs,
    BrowserFillFormArgs,
    BrowserHandleDialogArgs,
    BrowserHoverArgs,
    BrowserNavigateArgs,
    BrowserNavigateBackArgs,
    BrowserNetworkRequestsArgs,
    BrowserPressKeyArgs,
    BrowserResizeArgs,
    BrowserRunCodeArgs,
    BrowserSelectOptionArgs,
    BrowserSnapshotArgs,
    BrowserTabsArgs,
    BrowserTakeScreenshotArgs,
    BrowserTypeArgs,
    BrowserWaitForArgs,
    # Legacy/Internal
    NavigateArgs,
    GetPageStateArgs,
    PlaywrightTestArgs,
    ValidateTestArgs,
    AssertArgs,
    WaitArgs,
    ScrollArgs,
    KeyPressArgs,
    ReplaceArgs,
)

logger = logging.getLogger(__name__)


class BrowserTools:
    def __init__(self, page: Page, session_id: str = None):
        self.page = page
        self.session_id = session_id or str(uuid.uuid4())[:8]
        self.ref_map: Dict[str, str] = {}  # ref -> selector (or internal ID)
        self.console_logs: List[Dict[str, str]] = []
        self.network_logs: List[Dict[str, Any]] = []

        # Setup listeners
        self.page.on(
            "console",
            lambda msg: self.console_logs.append(
                {"type": msg.type, "text": msg.text, "location": msg.location}
            ),
        )
        self.page.on(
            "request",
            lambda req: self.network_logs.append(
                {
                    "method": req.method,
                    "url": req.url,
                    "resourceType": req.resource_type,
                }
            ),
        )

    async def _take_screenshot_base64(self) -> str:
        """Helper to take a screenshot and return base64 string."""
        try:
            screenshot_bytes = await self.page.screenshot(type="png")
            return base64.b64encode(screenshot_bytes).decode("utf-8")
        except Exception:
            return ""

    async def _error_with_screenshot(self, prefix: str, error: Exception) -> dict:
        """Helper to return an error dictionary with an optional screenshot."""
        screenshot = await self._take_screenshot_base64()
        return {
            "error": f"{prefix}: {str(error)}",
            "screenshot_base64": screenshot,
        }

    # --- MCP Style Tools ---

    async def browser_navigate(self, args: BrowserNavigateArgs) -> dict:
        try:
            await self.page.goto(args.url, wait_until="domcontentloaded")
            return {"message": f"Navigated to {args.url}"}
        except Exception as e:
            return await self._error_with_screenshot("Navigation failed", e)

    async def browser_navigate_back(self, args: BrowserNavigateBackArgs) -> dict:
        try:
            await self.page.go_back()
            return {"message": "Navigated back"}
        except Exception as e:
            return await self._error_with_screenshot("Go back failed", e)

    async def browser_click(self, args: BrowserClickArgs) -> dict:
        try:
            selector = self.ref_map.get(args.ref)
            if not selector:
                return {
                    "error": f"Invalid reference: {args.ref}. Element may no longer exist or snapshot is outdated.",
                    "screenshot_base64": await self._take_screenshot_base64(),
                }

            locator = self.page.locator(selector).first

            click_options = {
                "button": args.button,
                "modifiers": args.modifiers,
                "click_count": 2 if args.doubleClick else 1,
                "timeout": 15000,
            }

            await locator.click(**click_options)

            # After a click, wait a bit for potential navigation or dynamic updates
            # to prevent race conditions in subsequent snapshots.
            try:
                await self.page.wait_for_load_state("domcontentloaded", timeout=3000)
            except Exception:
                pass  # Not all clicks trigger navigation

            return {"message": f"Clicked element {args.ref}"}
        except Exception as e:
            return await self._error_with_screenshot("Click failed", e)

    async def browser_hover(self, args: BrowserHoverArgs) -> dict:
        try:
            selector = self.ref_map.get(args.ref)
            if not selector:
                return {
                    "error": f"Invalid reference: {args.ref}",
                    "screenshot_base64": await self._take_screenshot_base64(),
                }

            await self.page.locator(selector).first.hover(timeout=15000)
            return {"message": f"Hovered over element {args.ref}"}
        except Exception as e:
            return await self._error_with_screenshot("Hover failed", e)

    async def browser_type(self, args: BrowserTypeArgs) -> dict:
        try:
            selector = self.ref_map.get(args.ref)
            if not selector:
                return {
                    "error": f"Invalid reference: {args.ref}",
                    "screenshot_base64": await self._take_screenshot_base64(),
                }

            locator = self.page.locator(selector).first
            if args.slowly:
                await locator.press_sequentially(args.text, delay=100, timeout=15000)
            else:
                await locator.fill(args.text, timeout=15000)

            if args.submit:
                await self.page.keyboard.press("Enter")
                # Wait for potential navigation
                try:
                    await self.page.wait_for_load_state(
                        "domcontentloaded", timeout=3000
                    )
                except Exception:
                    pass

            return {"message": f"Typed into element {args.ref}"}
        except Exception as e:
            return await self._error_with_screenshot("Type failed", e)

    async def browser_fill_form(self, args: BrowserFillFormArgs) -> dict:
        results = []
        try:
            for field in args.fields:
                selector = self.ref_map.get(field.ref)
                if not selector:
                    results.append(
                        f"Field '{field.name}': Invalid reference {field.ref}"
                    )
                    continue

                locator = self.page.locator(selector).first

                if field.type == "textbox" or field.type == "slider":
                    await locator.fill(field.value, timeout=5000)
                elif field.type == "checkbox" or field.type == "radio":
                    if field.value.lower() == "true":
                        await locator.check(timeout=5000)
                    else:
                        await locator.uncheck(timeout=5000)
                elif field.type == "combobox":
                    await locator.select_option(field.value, timeout=5000)

                results.append(f"Field '{field.name}': Success")
                await asyncio.sleep(0.5)  # Small delay between fields for stability

            return {"message": "\n".join(results)}
        except Exception as e:
            res = await self._error_with_screenshot("Fill form failed", e)
            res["partial_results"] = results
            return res

    async def browser_select_option(self, args: BrowserSelectOptionArgs) -> dict:
        try:
            selector = self.ref_map.get(args.ref)
            if not selector:
                return {
                    "error": f"Invalid reference: {args.ref}",
                    "screenshot_base64": await self._take_screenshot_base64(),
                }

            await self.page.locator(selector).first.select_option(
                args.values, timeout=15000
            )
            return {"message": f"Selected options for element {args.ref}"}
        except Exception as e:
            return await self._error_with_screenshot("Select failed", e)

    async def browser_drag(self, args: BrowserDragArgs) -> dict:
        try:
            start_selector = self.ref_map.get(args.startRef)
            end_selector = self.ref_map.get(args.endRef)
            if not start_selector or not end_selector:
                return {
                    "error": "Invalid reference(s)",
                    "screenshot_base64": await self._take_screenshot_base64(),
                }

            await self.page.locator(start_selector).first.drag_to(
                self.page.locator(end_selector).first, timeout=15000
            )
            return {"message": f"Dragged {args.startRef} to {args.endRef}"}
        except Exception as e:
            return await self._error_with_screenshot("Drag failed", e)

    async def browser_evaluate(self, args: BrowserEvaluateArgs) -> dict:
        try:
            if args.ref:
                selector = self.ref_map.get(args.ref)
                if not selector:
                    return {
                        "error": f"Invalid reference: {args.ref}",
                        "screenshot_base64": await self._take_screenshot_base64(),
                    }
                element = await self.page.locator(selector).first.element_handle()
                result = await self.page.evaluate(args.function, element)
            else:
                result = await self.page.evaluate(args.function)
            return {"result": result}
        except Exception as e:
            return await self._error_with_screenshot("Evaluate failed", e)

    async def browser_run_code(self, args: BrowserRunCodeArgs) -> dict:
        # Dangerous but requested. We wrap it in an async function.
        try:
            # Note: This is a hacky way to run arbitrary code.
            # In a real system this would be more controlled.
            # We provide 'page' and 'expect' to the context.
            from playwright.async_api import expect

            # Create a namespace for execution
            namespace = {"page": self.page, "expect": expect, "asyncio": asyncio}

            # Wrap the code in an async function
            exec_code = "async def __temp_run():\n" + "\n".join(
                [f"    {line}" for line in args.code.split("\n")]
            )
            exec(exec_code, namespace)
            await namespace["__temp_run"]()

            return {"message": "Code executed successfully"}
        except Exception as e:
            return await self._error_with_screenshot("Code execution failed", e)

    async def browser_wait_for(self, args: BrowserWaitForArgs) -> dict:
        try:
            if args.time:
                await asyncio.sleep(args.time)
            if args.text:
                await self.page.get_by_text(args.text).wait_for(
                    state="visible", timeout=15000
                )
            if args.textGone:
                await self.page.get_by_text(args.textGone).wait_for(
                    state="hidden", timeout=15000
                )
            return {"message": "Wait complete"}
        except Exception as e:
            return await self._error_with_screenshot("Wait failed", e)

    async def browser_resize(self, args: BrowserResizeArgs) -> dict:
        try:
            await self.page.set_viewport_size(
                {"width": int(args.width), "height": int(args.height)}
            )
            return {"message": f"Resized to {args.width}x{args.height}"}
        except Exception as e:
            return await self._error_with_screenshot("Resize failed", e)

    async def browser_close(self, args: BrowserCloseArgs) -> dict:
        try:
            await self.page.close()
            return {"message": "Browser closed"}
        except Exception as e:
            # No screenshot for close failure as page might be gone
            return {"error": f"Close failed: {str(e)}"}

    async def browser_console_messages(self, args: BrowserConsoleMessagesArgs) -> dict:
        # Filters logs by level
        levels = ["error", "warning", "info", "debug"]
        target_idx = levels.index(args.level)
        relevant_levels = levels[: target_idx + 1]

        filtered = [log for log in self.console_logs if log["type"] in relevant_levels]

        if args.filename:
            path = os.path.join("logs", args.filename)
            os.makedirs("logs", exist_ok=True)
            with open(path, "w") as f:
                import json

                json.dump(filtered, f)
            return {"message": f"Logs saved to {path}"}

        return {"messages": filtered}

    async def browser_network_requests(self, args: BrowserNetworkRequestsArgs) -> dict:
        filtered = self.network_logs
        if not args.includeStatic:
            static_exts = (
                ".png",
                ".jpg",
                ".jpeg",
                ".gif",
                ".css",
                ".js",
                ".woff",
                ".woff2",
                ".svg",
            )
            filtered = [
                entry
                for entry in filtered
                if not any(entry["url"].lower().endswith(ext) for ext in static_exts)
            ]

        if args.filename:
            path = os.path.join("logs", args.filename)
            os.makedirs("logs", exist_ok=True)
            with open(path, "w") as f:
                import json

                json.dump(filtered, f)
            return {"message": f"Network logs saved to {path}"}

        return {"requests": filtered}

    async def browser_snapshot(self, args: BrowserSnapshotArgs) -> dict:
        # Returns current page state as a formatted snapshot
        state = await self.get_page_state(GetPageStateArgs())
        if args.filename:
            path = os.path.join("logs", args.filename)
            os.makedirs("logs", exist_ok=True)
            with open(path, "w") as f:
                f.write(state)
            return {"message": f"Snapshot saved to {path}"}
        return {"snapshot": state}

    async def browser_take_screenshot(self, args: BrowserTakeScreenshotArgs) -> dict:
        try:
            options = {"type": args.type, "full_page": args.fullPage}
            if args.ref:
                selector = self.ref_map.get(args.ref)
                if not selector:
                    return {"error": f"Invalid reference: {args.ref}"}
                screenshot_bytes = await self.page.locator(selector).first.screenshot(
                    **options
                )
            else:
                screenshot_bytes = await self.page.screenshot(**options)

            b64 = base64.b64encode(screenshot_bytes).decode("utf-8")

            if args.filename:
                path = os.path.join(
                    "videos", args.filename
                )  # Reuse videos dir for artifacts
                with open(path, "wb") as f:
                    f.write(screenshot_bytes)
                return {
                    "message": f"Screenshot saved to {path}",
                    "screenshot_base64": b64,
                }

            return {"message": "Screenshot taken", "screenshot_base64": b64}
        except Exception as e:
            return {"error": f"Screenshot failed: {str(e)}"}

    async def browser_handle_dialog(self, args: BrowserHandleDialogArgs) -> dict:
        # This is tricky because Playwright handles dialogs via events.
        # We'd need to have a pre-registered listener.
        # For now, we return a message.
        return {
            "message": "Dialog handling requested. Note: This tool currently only sets the next expected dialog action if implemented with a listener."
        }

    async def browser_tabs(self, args: BrowserTabsArgs) -> dict:
        try:
            pages = self.page.context.pages
            if args.action == "list":
                return {
                    "pages": [
                        {"index": i, "url": p.url, "title": await p.title()}
                        for i, p in enumerate(pages)
                    ]
                }
            elif args.action == "new":
                await self.page.context.new_page()
                return {"message": "New tab opened", "index": len(pages)}
            elif args.action == "select":
                if args.index is not None and 0 <= int(args.index) < len(pages):
                    self.page = pages[int(args.index)]
                    return {"message": f"Selected tab {args.index}"}
                return {"error": "Invalid index"}
            elif args.action == "close":
                idx = (
                    int(args.index)
                    if args.index is not None
                    else pages.index(self.page)
                )
                await pages[idx].close()
                return {"message": f"Closed tab {idx}"}
        except Exception as e:
            return {"error": f"Tabs action failed: {str(e)}"}

    async def browser_install(self) -> dict:
        return {"message": "Browser is already installed in this environment."}

    async def browser_press_key(self, args: BrowserPressKeyArgs) -> dict:
        try:
            await self.page.keyboard.press(args.key)
            return {"message": f"Pressed key: {args.key}"}
        except Exception as e:
            return await self._error_with_screenshot("Press key failed", e)

    async def browser_file_upload(self, args: BrowserFileUploadArgs) -> dict:
        """
        Upload files by setting them directly on file input elements.
        This avoids the file chooser dialog entirely.
        """
        try:
            if not args.paths:
                return {"message": "File chooser cancelled (no paths provided)"}

            # Validate that all files exist
            for path in args.paths:
                if not os.path.exists(path):
                    return {
                        "error": f"File not found: {path}",
                        "screenshot_base64": await self._take_screenshot_base64(),
                    }

            # Find all file input elements on the page
            file_inputs = await self.page.locator('input[type="file"]').all()

            if not file_inputs:
                return {
                    "error": "No file input elements found on the page",
                    "screenshot_base64": await self._take_screenshot_base64(),
                }

            # Set files on the first visible file input
            # (or the first one if none are visible)
            uploaded = False
            for file_input in file_inputs:
                try:
                    # Set files directly on the input element
                    await file_input.set_input_files(args.paths)
                    uploaded = True
                    logger.info(f"Successfully uploaded files: {args.paths}")
                    break
                except Exception as e:
                    logger.debug(f"Failed to upload to this input: {e}")
                    continue

            if not uploaded:
                return {
                    "error": "Failed to upload files to any file input element",
                    "screenshot_base64": await self._take_screenshot_base64(),
                }

            return {
                "message": f"Successfully uploaded {len(args.paths)} file(s): {', '.join(os.path.basename(p) for p in args.paths)}"
            }

        except Exception as e:
            return await self._error_with_screenshot("File upload failed", e)

    # --- Legacy/FlowTest Specific Implementation (Updated to be more robust) ---

    async def playwright_test(self, args: PlaywrightTestArgs) -> str:
        return args.code

    async def validate_test(self, args: ValidateTestArgs) -> str:
        try:
            if not os.path.exists("tests_js"):
                os.makedirs("tests_js")

            spec_content = f"""import {{ test, expect }} from '@playwright/test';
test('validation run', async ({{ page }}) => {{
    {args.code}
}});"""
            spec_name = f"validate_{self.session_id}.spec.js"
            spec_path = os.path.join("tests_js", spec_name)
            with open(spec_path, "w") as f:
                f.write(spec_content)

            result = subprocess.run(
                ["npx", "playwright", "test", spec_name, "--timeout=15000"],
                cwd="tests_js",
                capture_output=True,
                text=True,
            )

            if os.path.exists(spec_path):
                os.remove(spec_path)

            if result.returncode == 0:
                # Return the code itself on success - this auto-submits the validated code
                return args.code
            else:
                return f"VALIDATION FAILED:\n{result.stdout}\n{result.stderr}"
        except Exception as e:
            return f"Validation tool error: {str(e)}"

    async def navigate(self, args: NavigateArgs) -> dict:
        return await self.browser_navigate(BrowserNavigateArgs(url=args.url))

    async def get_page_state(self, args: GetPageStateArgs) -> str:
        """
        Enhanced version that assigns 'refs' to elements and returns a snapshot.
        """
        try:
            # Wait for any pending navigation or dynamic updates to settle.
            # This prevents the agent from seeing stale states during transitions.
            try:
                # Use a combination of load states for maximum stability
                await self.page.wait_for_load_state("domcontentloaded", timeout=2000)
                await asyncio.sleep(0.5)  # Mandatory settle time
            except Exception:
                pass

            title = await self.page.title()
            url = self.page.url

            # Use aria_snapshot for the LLM to read
            snapshot = await self.page.locator("body").aria_snapshot()

            # Injecting a script to get interactive elements and their current "state"
            # and mapping them to a unique ref.

            # Clean ref map for new state
            self.ref_map = {}

            # This script finds all potentially interactive elements and returns their
            # accessible names, roles, and a way to reach them.
            # Note: We don't actually need the complex tree in self.ref_map,
            # we just need to be able to find the element again.
            # Using data-flowtest-ref attribute is the most reliable.

            try:
                # 1. Clean up old refs and tag new ones in one go
                # 2. Filter for visibility to avoid tagging hidden inputs/elements
                mapping_data = await self.page.evaluate("""
                    () => {
                        // Clear existing tags to ensure no duplicates
                        document.querySelectorAll('[data-flowtest-ref]').forEach(el => el.removeAttribute('data-flowtest-ref'));

                        let count = 0;
                        const selectors = [
                            'button', 'input:not([type="hidden"])', 'select', 'textarea', 'a',
                            '[role="button"]', '[role="link"]', '[role="checkbox"]', '[role="menuitem"]', '[role="option"]', '[role="combobox"]', '[role="listbox"]',
                            '[aria-haspopup]', '[aria-expanded]', '[onclick]',
                            'h1', 'h2', 'h3', 'p', 'span', 'label'
                        ].join(',');

                        const elements = Array.from(document.querySelectorAll(selectors)).filter(el => {
                            const style = window.getComputedStyle(el);
                            return style.display !== 'none' &&
                                   style.visibility !== 'hidden' &&
                                   el.offsetWidth > 0 &&
                                   el.offsetHeight > 0;
                        });

                        // ENHANCEMENT: Also capture generic divs/spans with cursor: pointer
                        // This catches React/Vue components that use JavaScript event listeners
                        const cursorPointerElements = Array.from(
                            document.querySelectorAll('div[class], span[class]')
                        ).filter(el => {
                            const style = window.getComputedStyle(el);
                            const text = (el.innerText || '').trim();
                            return style.cursor === 'pointer' &&
                                   style.display !== 'none' &&
                                   style.visibility !== 'hidden' &&
                                   el.offsetWidth > 0 &&
                                   el.offsetHeight > 0 &&
                                   text.length > 0 && // Has text content
                                   text.length < 300 && // Not too long (avoid page containers)
                                   !el.hasAttribute('data-flowtest-ref'); // Not already tagged
                        });

                        // Merge both lists
                        const allInteractiveElements = [...elements, ...cursorPointerElements];

                        return allInteractiveElements.map(el => {
                            const ref = String(++count);
                            el.setAttribute('data-flowtest-ref', ref);
                            return {
                                ref: ref,
                                role: el.role || el.tagName.toLowerCase(),
                                name: (el.innerText || el.ariaLabel || el.placeholder || el.value || "").trim().substring(0, 80)
                            };
                        });
                    }
                """)

                # Populate ref_map correctly
                self.ref_map = {
                    item["ref"]: f'[data-flowtest-ref="{item["ref"]}"]'
                    for item in mapping_data
                }
            except Exception as e:
                logger.error(f"Error tagging elements: {e}")
                mapping_data = []

            ref_list = "\n".join(
                [
                    f'- {item["role"]} "{item["name"]}" [ref="{item["ref"]}"]'
                    for item in mapping_data
                ]
            )

            output = dedent(
                f"""
                --- PAGE STATE ---
                Title: {title}
                URL: {url}

                --- INTERACTIVE & CONTENT ELEMENTS (with refs) ---
                {ref_list}

                --- ACCESSIBILITY SNAPSHOT ---
                {snapshot}

                --- GUIDANCE ---
                Use the 'ref' from the list above for your tool calls (click, type, assert).
                Example: assert_state(ref="5", condition="contains_text", value="Success")
            """.strip()
            )
            return output
        except Exception as e:
            return f"Failed to get page state: {str(e)}"

    async def assert_state(self, args: AssertArgs) -> str:
        try:
            selector = self.ref_map.get(args.ref) if args.ref else args.selector
            if not selector:
                return "Assertion failed: No 'ref' or 'selector' provided."

            locator = self.page.locator(selector).first
            if args.condition == "visible":
                if await locator.is_visible(timeout=15000):
                    return "ASSERTION PASSED: Element is visible."
                else:
                    return "ASSERTION FAILED: Element is NOT visible."
            elif args.condition == "hidden":
                if await locator.is_hidden(timeout=15000):
                    return "ASSERTION PASSED: Element is hidden."
                else:
                    return "ASSERTION FAILED: Element is visible."
            elif args.condition == "contains_text":
                text = await locator.text_content(timeout=15000)
                if args.value in (text or ""):
                    return f"ASSERTION PASSED: Element contains '{args.value}'."
                else:
                    return f"ASSERTION FAILED: Expected element to contain '{args.value}', found '{text}'."
            elif args.condition == "url_contains":
                if args.value in self.page.url:
                    return f"ASSERTION PASSED: URL contains '{args.value}'."
                else:
                    return f"ASSERTION FAILED: URL '{self.page.url}' does NOT contain '{args.value}'."
            return "Unknown assertion condition."
        except Exception as e:
            return f"Assertion error: {str(e)}"

    async def scroll(self, args: ScrollArgs) -> str:
        # Implementation remains similar as it doesn't use refs
        try:
            if args.direction == "top":
                await self.page.evaluate("window.scrollTo(0, 0)")
            elif args.direction == "bottom":
                await self.page.evaluate(
                    "window.scrollTo(0, document.body.scrollHeight)"
                )
            elif args.direction == "up":
                await self.page.mouse.wheel(0, -args.amount)
            else:  # down
                await self.page.mouse.wheel(0, args.amount)
            await asyncio.sleep(0.5)
            return f"Scrolled {args.direction}"
        except Exception as e:
            return f"Scroll failed: {str(e)}"

    async def press_key(self, args: KeyPressArgs) -> str:
        return (await self.browser_press_key(BrowserPressKeyArgs(key=args.key)))[
            "message"
        ]

    async def wait(self, args: WaitArgs) -> str:
        await asyncio.sleep(min(args.seconds, 5.0))
        return f"Waited {args.seconds}s"


class DocumentTools:
    def __init__(self, file: str):
        self.file = file
        self.document = Document(file)

    async def replace(self, args: ReplaceArgs) -> str:
        count = 0

        def _replace_in_paragraph(paragraph):
            """Replace text while preserving run formatting"""
            if args.search_text not in paragraph.text:
                return False

            full_text = paragraph.text
            new_text = full_text.replace(args.search_text, args.replace_text)

            # Preserve formatting by keeping first run's style
            if len(paragraph.runs) > 0:
                first_run = paragraph.runs[0]
                # Store formatting
                font_name = first_run.font.name
                font_size = first_run.font.size
                font_bold = first_run.font.bold
                font_italic = first_run.font.italic
                font_color = (
                    first_run.font.color.rgb if first_run.font.color.rgb else None
                )

                # Clear runs
                for run in paragraph.runs:
                    run._element.getparent().remove(run._element)

                # Add back with same formatting
                new_run = paragraph.add_run(new_text)
                new_run.font.name = font_name
                new_run.font.size = font_size
                new_run.font.bold = font_bold
                new_run.font.italic = font_italic
                if font_color:
                    new_run.font.color.rgb = font_color
            else:
                paragraph.add_run(new_text)

            return True

        # Loop through paragraphs to replace text
        for paragraph in self.document.paragraphs:
            if _replace_in_paragraph(paragraph):
                count += 1
                logger.debug("search_text found!")

        # Also check tables (optional, but recommended)
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if _replace_in_paragraph(paragraph):
                            count += 1
                            logger.debug("search_text found!")

        if count == 1:
            self.document.save(self.file)
            logger.debug(f"Saved as {self.file}")
            return "Successfully replaced"
        elif count == 0:
            return "ERROR: search_text not found in document. Make sure to include exact text from the resume including newlines and spacing. Consider copying-pasting directly from the resume."
        else:
            return f"ERROR: search_text appears {count} times in the resume. To fix this, include MORE CONTEXT (dates, section headers, job titles, adjacent bullets) to make your search_text unique and appear only ONCE. Example: instead of searching just the bullet point, include the job title and date before it."
